import base64
from pathlib import Path

from docx import Document

from report_backend.parsers.markdown_renderer import render_markdown_to_docx


PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Wnki0sAAAAASUVORK5CYII="
)


def test_render_markdown_to_docx_handles_tables_and_missing_images(tmp_path: Path) -> None:
    image_path = tmp_path / "sample.png"
    image_path.write_bytes(PNG_BYTES)
    output_path = tmp_path / "report.docx"

    markdown = "\n".join(
        [
            "# Title",
            "",
            "This is a paragraph.",
            "",
            f"![Sample Image]({image_path.as_posix()})",
            "",
            "![Broken Image](missing.png)",
            "",
            "| Item | Content |",
            "| --- | --- |",
            "| A | B |",
        ]
    )
    render_markdown_to_docx(markdown, output_path)

    document = Document(output_path)
    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert output_path.exists()
    assert "Title" in full_text
    assert "missing.png" in full_text
    assert len(document.inline_shapes) == 1
    assert len(document.tables) == 1


def test_render_markdown_to_docx_tolerates_ragged_tables(tmp_path: Path) -> None:
    output_path = tmp_path / "ragged-table.docx"
    markdown = "\n".join(
        [
            "# Table Stress",
            "",
            "| Col A | Col B | Col C |",
            "| --- | --- | --- |",
            "| 1 | 2 |",
            "| 3 | 4 | 5 | 6 |",
        ]
    )

    render_markdown_to_docx(markdown, output_path)

    document = Document(output_path)
    assert output_path.exists()
    assert len(document.tables) == 1
    assert len(document.tables[0].rows) == 3
    assert len(document.tables[0].columns) == 4
