from pathlib import Path

from docx import Document

from report_backend.parsers.template_parser import extract_template_outline


def test_extract_template_outline_from_example(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    document = Document()
    document.add_heading("一、实验目的", level=1)
    document.add_paragraph("完成系统分析。")
    document.add_heading("二、实验步骤", level=1)
    document.save(template_path)

    outline = extract_template_outline(template_path)

    assert outline
    assert "# 一、实验目的" in outline
    assert "# 二、实验步骤" in outline
