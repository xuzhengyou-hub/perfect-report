import io
import json
import time
from pathlib import Path

import fitz
import pytest
from docx import Document
from fastapi.testclient import TestClient

from report_backend.api import dependencies
from report_backend.main import app


def _build_template(path: Path) -> None:
    document = Document()

    cover = document.add_table(rows=6, cols=2)
    cover.cell(0, 0).text = "课程名称："
    cover.cell(1, 0).text = "实验项目："
    cover.cell(2, 0).text = "实验时间："
    cover.cell(3, 0).text = "实验班级："
    cover.cell(4, 0).text = "总 份 数"
    cover.cell(5, 0).text = "指导教师："
    cover.cell(0, 1).text = "软件工程"
    cover.cell(2, 1).text = "2025/12/16"
    cover.cell(4, 1).text = "1份"
    cover.cell(5, 1).text = "黄超"

    roster = document.add_table(rows=3, cols=15)
    roster.cell(0, 0).text = "学院："
    roster.cell(0, 4).text = "专业："
    roster.cell(0, 9).text = "班级："
    roster.cell(0, 13).text = "成绩："
    roster.cell(1, 0).text = "姓名："
    roster.cell(1, 3).text = "学号："
    roster.cell(1, 8).text = "组别："
    roster.cell(1, 12).text = "组员："
    roster.cell(2, 0).text = "实验地点："
    roster.cell(2, 5).text = "实验日期："
    roster.cell(2, 11).text = "指导教师签名："

    evaluation = document.add_table(rows=1, cols=8)
    evaluation.cell(0, 0).text = "预习情况"
    evaluation.cell(0, 2).text = "操作情况"
    evaluation.cell(0, 4).text = "考勤情况"
    evaluation.cell(0, 6).text = "数据处理情况"

    document.add_paragraph("实验       项目名称：")
    document.add_paragraph("实验目的：")
    document.add_paragraph("旧的实验目的占位内容")
    document.add_paragraph("实验内容：")
    document.add_paragraph("旧的实验内容占位内容")
    document.add_paragraph("实验步骤：")
    document.add_paragraph("旧的实验步骤占位内容")
    document.add_paragraph("序列图组成：")
    document.add_paragraph("旧的序列图组成占位内容")
    document.add_paragraph("协作图组成：")
    document.add_paragraph("旧的协作图组成占位内容")
    document.add_paragraph("分析飞机机票的事件流：")
    document.add_paragraph("旧的事件流占位内容")
    document.add_paragraph("创建飞机订票系统序列图的基本思路(通过图来表示)：")
    document.add_paragraph("旧的序列图思路占位内容")
    document.add_paragraph("创建飞机订票系统协作图的操作步骤(通过图来表示)：")
    document.add_paragraph("旧的协作图步骤占位内容")
    document.add_paragraph("实验小结：")
    document.add_paragraph("旧的小结占位内容")
    document.save(path)


def _read_success_event(client: TestClient, task_id: str) -> dict[str, object]:
    with client.stream("GET", f"/api/v1/report/events/{task_id}") as response:
        assert response.status_code == 200
        deadline = time.time() + 20
        for line in response.iter_lines():
            if not line or not line.startswith("data: "):
                if time.time() > deadline:
                    break
                continue
            event = json.loads(line[6:])
            if event["status"] in {"SUCCESS", "FAILED"}:
                return event
    raise AssertionError("Did not receive terminal event in time")


@pytest.fixture(autouse=True)
def disable_dashscope(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(dependencies.adapter.settings, "dashscope_api_key", None)


def test_generate_report_end_to_end(tmp_path) -> None:
    client = TestClient(app)
    template_path = tmp_path / "template.docx"
    _build_template(template_path)
    template_bytes = template_path.read_bytes()
    text_bytes = "系统需要支持订票、退票和查询航班信息。".encode("utf-8")

    response = client.post(
        "/api/v1/report/generate",
        data={"title": "飞机订票系统实验", "extra_info": "需要重点说明 UML 图和系统交互。"},
        files=[
            (
                "template_file",
                ("实验报告模板.docx", io.BytesIO(template_bytes), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            ),
            ("background_files", ("资料.txt", io.BytesIO(text_bytes), "text/plain")),
        ],
    )

    assert response.status_code == 200
    task_id = response.json()["data"]["task_id"]

    event = _read_success_event(client, task_id)
    assert event["status"] == "SUCCESS"
    assert event["file_name"] == "result_report.docx"

    workspace = dependencies.workspace_manager.get_workspace(task_id)
    synthesized_markdown_path = workspace / "outputs" / "synthesized_knowledge.md"
    final_markdown_path = workspace / "outputs" / "final_report.md"
    structured_json_path = workspace / "outputs" / "report_structured.json"
    assert synthesized_markdown_path.exists()
    assert final_markdown_path.exists()
    assert structured_json_path.exists()

    synthesized_markdown = synthesized_markdown_path.read_text(encoding="utf-8")
    final_markdown = final_markdown_path.read_text(encoding="utf-8")
    assert synthesized_markdown.startswith("# Synthesized Knowledge: ")
    assert "Experiment title:" in synthesized_markdown
    assert final_markdown.startswith("#")
    assert "## 模板字段" in final_markdown

    output_path = workspace / "outputs" / "result_report.docx"
    document = Document(output_path)
    assert len(document.tables) == 3
    assert document.tables[0].cell(1, 1).text == "飞机订票系统实验"
    assert "【待用户填写" in document.tables[0].cell(3, 1).text

    download = client.get(f"/api/v1/report/download/{task_id}")
    assert download.status_code == 200
    assert download.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def test_generate_report_preserves_warning_for_unsupported_background_file(tmp_path) -> None:
    client = TestClient(app)
    template_path = tmp_path / "template.docx"
    _build_template(template_path)

    response = client.post(
        "/api/v1/report/generate",
        data={"title": "Warning Check", "extra_info": ""},
        files=[
            (
                "template_file",
                (
                    "template.docx",
                    io.BytesIO(template_path.read_bytes()),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            ),
            ("background_files", ("notes.md", io.BytesIO(b"# unsupported"), "text/markdown")),
        ],
    )

    assert response.status_code == 200
    task_id = response.json()["data"]["task_id"]

    event = _read_success_event(client, task_id)
    assert event["status"] == "SUCCESS"
    assert event["warning"]
    assert "notes.md" in event["warning"]


def test_generate_report_writes_parsed_markdown_for_pdf_images(tmp_path) -> None:
    client = TestClient(app)
    template_path = tmp_path / "template.docx"
    _build_template(template_path)

    pdf_path = tmp_path / "background.pdf"
    pdf_document = fitz.open()
    page = pdf_document.new_page()
    page.insert_text((72, 72), "Signal flow diagram.")
    page.insert_image(
        fitz.Rect(72, 120, 180, 220),
        stream=b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x04\x00\x00\x00\xb5\x1c\x0c\x02\x00\x00\x00\x0bIDATx\xdac\xfc\xff\x1f\x00\x03\x03\x02\x00\xee\x95\xa9\xea\x00\x00\x00\x00IEND\xaeB`\x82",
    )
    pdf_document.save(pdf_path)
    pdf_document.close()

    response = client.post(
        "/api/v1/report/generate",
        data={"title": "PDF Image Flow", "extra_info": ""},
        files=[
            (
                "template_file",
                (
                    "template.docx",
                    io.BytesIO(template_path.read_bytes()),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            ),
            ("background_files", ("background.pdf", io.BytesIO(pdf_path.read_bytes()), "application/pdf")),
        ],
    )

    assert response.status_code == 200
    task_id = response.json()["data"]["task_id"]

    event = _read_success_event(client, task_id)
    assert event["status"] == "SUCCESS"
    assert "Image understanding skipped" in event["warning"]

    workspace = dependencies.workspace_manager.get_workspace(task_id)
    parsed_markdown_path = workspace / "outputs" / "parsed_background.md"
    synthesized_markdown_path = workspace / "outputs" / "synthesized_knowledge.md"
    final_markdown_path = workspace / "outputs" / "final_report.md"
    structured_json_path = workspace / "outputs" / "report_structured.json"
    assert parsed_markdown_path.exists()
    assert synthesized_markdown_path.exists()
    assert final_markdown_path.exists()
    assert structured_json_path.exists()

    parsed_markdown = parsed_markdown_path.read_text(encoding="utf-8")
    synthesized_markdown = synthesized_markdown_path.read_text(encoding="utf-8")
    final_markdown = final_markdown_path.read_text(encoding="utf-8")
    assert "# PDF Material: background" in parsed_markdown
    assert "Signal flow diagram." in parsed_markdown
    assert "![" in parsed_markdown
    assert "# Synthesized Knowledge: PDF Image Flow" in synthesized_markdown
    assert "Signal flow diagram." in synthesized_markdown
    assert "Reference Images" in synthesized_markdown
    assert "## 图片槽位" in final_markdown

    output_path = workspace / "outputs" / "result_report.docx"
    document = Document(output_path)
    assert len(document.inline_shapes) >= 1
