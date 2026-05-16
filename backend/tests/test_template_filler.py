from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from report_backend.parsers.template_filler import (
    TemplateReport,
    finalize_template_report,
    render_template_report,
)


def _tiny_png_bytes() -> bytes:
    return (
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
        b"\x08\x04\x00\x00\x00\xb5\x1c\x0c\x02\x00\x00\x00\x0bIDATx\xdac\xfc"
        b"\xff\x1f\x00\x03\x03\x02\x00\xee\x95\xa9\xea\x00\x00\x00\x00IEND\xaeB`\x82"
    )


def _build_detailed_template(path: Path) -> None:
    document = Document()

    cover = document.add_table(rows=6, cols=2)
    cover.cell(0, 0).text = "课程名称："
    cover.cell(1, 0).text = "实验项目："
    cover.cell(2, 0).text = "实验时间："
    cover.cell(3, 0).text = "实验班级："
    cover.cell(4, 0).text = "总份数："
    cover.cell(5, 0).text = "指导教师："

    cover_value = cover.cell(0, 1).paragraphs[0]
    cover_value.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_value.add_run("软件工程")
    cover_value.runs[0].font.size = Pt(16)
    cover_value.runs[0].bold = True

    cover.cell(2, 1).text = "2025/12/16"
    cover.cell(4, 1).text = "1份"
    cover.cell(5, 1).text = "黄超"

    roster = document.add_table(rows=3, cols=15)
    roster.cell(0, 0).text = "学院："
    roster.cell(0, 4).text = "专业："
    roster.cell(0, 9).text = "班级："
    roster.cell(0, 13).text = "成绩："
    roster.cell(1, 0).text = "姓名："
    roster.cell(1, 3).text = "学号："
    roster.cell(1, 8).text = "组别："
    roster.cell(1, 12).text = "组员："
    roster.cell(2, 0).text = "实验地点："
    roster.cell(2, 5).text = "实验日期："
    roster.cell(2, 11).text = "指导教师签名："

    roster_value = roster.cell(1, 1).paragraphs[0]
    roster_value.add_run("模板原值")
    roster_value.runs[0].font.size = Pt(12)
    roster_value.runs[0].italic = True

    evaluation = document.add_table(rows=1, cols=8)
    evaluation.cell(0, 0).text = "预习情况"
    evaluation.cell(0, 2).text = "操作情况"
    evaluation.cell(0, 4).text = "考勤情况"
    evaluation.cell(0, 6).text = "数据处理情况"

    title_paragraph = document.add_paragraph()
    title_run = title_paragraph.add_run("实验       项目名称：")
    title_run.font.size = Pt(15)
    title_run.bold = True

    def add_heading(text: str) -> None:
        p = document.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)

    def add_body(text: str) -> None:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.font.size = Pt(12)

    add_heading("实验目的：")
    add_body("旧的实验目的占位内容")
    add_heading("实验内容：")
    add_body("旧的实验内容占位内容")
    add_heading("实验步骤：")
    add_body("旧的实验步骤占位内容")
    add_heading("序列图组成：")
    add_body("旧的序列图组成占位内容")
    add_heading("协作图组成：")
    add_body("旧的协作图组成占位内容")
    add_heading("分析飞机订票的事件流：")
    add_body("旧的事件流占位内容")
    add_heading("创建飞机订票系统序列图的基本思路(通过图来表示)：")
    add_body("旧的序列图思路占位内容")
    add_heading("创建飞机订票系统协作图的操作步骤(通过图来表示)：")
    add_body("旧的协作图步骤占位内容")
    add_heading("实验小结：")
    add_body("旧的小结占位内容")

    document.save(path)


def _build_compact_template(path: Path) -> None:
    document = Document()
    cover = document.add_table(rows=6, cols=2)
    for row, label in enumerate(("课程名称：", "实验项目：", "实验时间：", "实验班级：", "总份数：", "指导教师：")):
        cover.cell(row, 0).text = label
    roster = document.add_table(rows=3, cols=15)
    evaluation = document.add_table(rows=1, cols=8)
    roster.cell(0, 0).text = "学院："
    evaluation.cell(0, 0).text = "预习情况"

    title = document.add_paragraph("实验 8 项目名称：")
    title.runs[0].bold = True

    def add_heading(text: str) -> None:
        p = document.add_paragraph(text)
        p.runs[0].bold = True

    def add_body(text: str) -> None:
        p = document.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.runs[0].font.size = Pt(12)

    add_heading("实验目的和要求")
    add_body("旧的目的内容")
    add_heading("实验内容及步骤")
    add_body("旧的内容内容")
    add_heading("实验结果与分析")
    add_body("旧的结果内容")
    add_heading("问题与讨论")
    add_body("旧的讨论内容")

    document.save(path)


def test_render_template_report_fills_known_fields_and_placeholders(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    output_path = tmp_path / "result.docx"
    image_path = tmp_path / "sequence.png"
    image_path.write_bytes(_tiny_png_bytes())
    _build_detailed_template(template_path)

    report = TemplateReport(
        cover_fields={
            "course_name": "软件工程",
            "experiment_name": "UML序列图和协作图实验",
            "experiment_time": "2025/12/16",
            "class_name": None,
            "total_score": None,
            "teacher_name": "黄超",
        },
        roster_fields={
            "college": "计算机科学学院",
            "major": "计算机科学与技术",
            "class_name": None,
            "score": None,
            "student_name": None,
            "student_id": None,
            "group_name": None,
            "group_members": None,
            "location": None,
            "experiment_date": "2025/12/16",
            "teacher_signature": None,
        },
        evaluation_fields={
            "preview_status": None,
            "operation_status": None,
            "attendance_status": None,
            "data_processing_status": None,
        },
        sections={
            "objective": ["掌握序列图和协作图的基本概念。"],
            "content": ["围绕飞机订票系统完成建模实验。"],
            "steps": ["分析事件流。", "绘制相关图示。"],
            "sequence_components": ["参与者、生命线、消息。"],
            "collaboration_components": ["对象、连接、消息顺序。"],
            "event_flow": ["用户登录。", "查询航班。"],
            "sequence_idea": ["根据事件流确定交互时序。"],
            "collaboration_steps": ["根据对象关系标注消息编号。"],
            "summary": ["已完成模板化填充。"],
        },
        images={
            "sequence_diagram": str(image_path),
            "collaboration_diagram": str(image_path),
        },
        gallery_images=[str(image_path)],
    )

    render_template_report(template_path, output_path, report)

    document = Document(output_path)
    assert len(document.tables) == 3
    assert document.tables[0].cell(1, 1).text == "UML序列图和协作图实验"
    assert "【待用户填写" in document.tables[0].cell(3, 1).text
    assert document.tables[1].cell(1, 1).text == "模板原值"
    assert any("项目名称" in paragraph.text and "UML序列图和协作图实验" in paragraph.text for paragraph in document.paragraphs)
    assert any("掌握序列图和协作图的基本概念。" in paragraph.text for paragraph in document.paragraphs)
    assert any("已完成模板化填充。" in paragraph.text for paragraph in document.paragraphs)
    assert len(document.inline_shapes) == 2

    with ZipFile(output_path) as archive:
        media_files = [name for name in archive.namelist() if name.startswith("word/media/")]
    assert media_files


def test_render_template_report_preserves_cell_and_body_formatting(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    output_path = tmp_path / "result.docx"
    _build_detailed_template(template_path)

    report = TemplateReport(
        cover_fields={
            "course_name": "高级软件工程",
            "experiment_name": "格式校验实验",
            "experiment_time": "2025/12/16",
            "class_name": None,
            "total_score": None,
            "teacher_name": "黄超",
        },
        roster_fields={
            "college": "计算机科学学院",
            "major": "计算机科学与技术",
            "class_name": None,
            "score": None,
            "student_name": "张三",
            "student_id": None,
            "group_name": None,
            "group_members": None,
            "location": None,
            "experiment_date": "2025/12/16",
            "teacher_signature": None,
        },
        evaluation_fields={},
        sections={
            "objective": ["第一段正文。", "第二段正文。"],
            "content": ["实验内容正文。"],
            "steps": ["步骤一。"],
            "sequence_components": ["组成一。"],
            "collaboration_components": ["组成二。"],
            "event_flow": ["事件流一。"],
            "sequence_idea": ["序列图说明。"],
            "collaboration_steps": ["协作图说明。"],
            "summary": ["总结。"],
        },
        images={},
        gallery_images=[],
    )

    render_template_report(template_path, output_path, report)
    document = Document(output_path)

    cover_paragraph = document.tables[0].cell(0, 1).paragraphs[0]
    cover_run = cover_paragraph.runs[0]
    assert cover_paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert cover_run.bold is True
    assert cover_run.font.size == Pt(16)

    roster_run = document.tables[1].cell(1, 1).paragraphs[0].runs[0]
    assert roster_run.italic is True
    assert roster_run.font.size == Pt(12)

    objective_paragraph = next(paragraph for paragraph in document.paragraphs if paragraph.text == "第一段正文。")
    objective_run = objective_paragraph.runs[0]
    assert objective_paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert objective_run.font.size == Pt(12)

    title_paragraph = next(paragraph for paragraph in document.paragraphs if "项目名称" in paragraph.text)
    title_run = title_paragraph.runs[0]
    assert title_run.bold is True
    assert title_run.font.size == Pt(15)


def test_finalize_and_render_compact_template_keeps_detailed_steps_and_all_images(tmp_path: Path) -> None:
    template_path = tmp_path / "compact.docx"
    output_path = tmp_path / "compact_result.docx"
    _build_compact_template(template_path)

    image_paths = []
    for index in range(3):
        image_path = tmp_path / f"image_{index}.png"
        image_path.write_bytes(_tiny_png_bytes())
        image_paths.append(str(image_path))

    report = TemplateReport(
        cover_fields={
            "course_name": "智能系统设计",
            "experiment_name": "语音识别与翻译实验",
            "experiment_time": "2026/05/17",
            "class_name": "计科2班",
            "total_score": None,
            "teacher_name": "温清木",
        },
        roster_fields={},
        evaluation_fields={},
        sections={
            "objective": ["掌握图形化语音识别基本方法。"],
            "content": ["使用行空板完成语音识别与翻译。"],
            "steps": ["旧步骤摘要。"],
            "sequence_components": [],
            "collaboration_components": [],
            "event_flow": [],
            "sequence_idea": [],
            "collaboration_steps": [],
            "summary": ["成功实现语音识别和中英翻译功能。"],
        },
        images={"sequence_diagram": image_paths[0], "collaboration_diagram": None},
        gallery_images=[],
    )

    synthesized_markdown = """
#### 实验内容及步骤
1. 硬件搭建：
   - 使用USB连接线将行空板连接到计算机。
2. 软件准备：
   - 打开Mind+编程软件，完成软件准备过程。
3. 编写程序：
   - **初始化块**：设置“讯飞语音初始化”。
   - **账号设置**：配置百度翻译API。
   - **显示对象**：创建中文和英文文本框。
   - **循环执行**：按键A录音识别，按键B翻译显示。

#### 实验结果与分析
成功实现语音识别和中英翻译功能。
"""

    finalized = finalize_template_report(
        report,
        synthesized_markdown=synthesized_markdown,
        parsed_markdown="",
        image_refs=image_paths,
    )
    render_template_report(template_path, output_path, finalized)

    document = Document(output_path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    assert "硬件搭建" in paragraphs
    assert any("初始化块" in paragraph and "讯飞语音初始化" in paragraph for paragraph in paragraphs)
    assert any("循环执行" in paragraph and "按键A录音识别" in paragraph for paragraph in paragraphs)
    assert len(document.inline_shapes) == 3
