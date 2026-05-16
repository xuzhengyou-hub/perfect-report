from __future__ import annotations

import re
from pathlib import Path

from docx import Document


HEADING_RE = re.compile(r"^[一二三四五六七八九十]+、")


def extract_template_outline(template_path: Path) -> str:
    document = Document(template_path)
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
        if style_name.startswith("heading 1") or HEADING_RE.match(text):
            lines.append(f"# {text}")
        elif style_name.startswith("heading 2"):
            lines.append(f"## {text}")
        else:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() or "[待用户补充: 表格内容]" for cell in row.cells]
            if any(values):
                lines.append("| " + " | ".join(values) + " |")

    return "\n".join(lines).strip()
