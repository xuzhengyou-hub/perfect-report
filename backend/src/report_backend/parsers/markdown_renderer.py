from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


IMAGE_RE = re.compile(r"!\[(?P<alt>.*?)\]\((?P<path>.*?)\)")


def render_markdown_to_docx(markdown: str, output_path: Path) -> Path:
    document = Document()
    _apply_default_styles(document)

    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped.lstrip("#").strip()
            document.add_heading(text, level=min(level, 4))
            index += 1
            continue
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = [stripped]
            index += 1
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            _add_table(document, table_lines)
            continue
        if IMAGE_RE.fullmatch(stripped):
            _add_image_or_fallback(document, stripped)
            index += 1
            continue
        paragraph = document.add_paragraph()
        _append_runs(paragraph, stripped)
        index += 1

    document.save(output_path)
    return output_path


def _apply_default_styles(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.25
    paragraph_format.space_after = Pt(6)


def _append_runs(paragraph, text: str) -> None:
    image_match = IMAGE_RE.search(text)
    if image_match and image_match.group(0) == text:
        _add_image_or_fallback(paragraph._parent, text)
        return
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _add_table(document: Document, table_lines: list[str]) -> None:
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if _is_separator_row(cells):
            continue
        rows.append(cells)
    if not rows:
        return
    max_cols = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    for row_index, row_values in enumerate(rows):
        padded_values = row_values + [""] * (max_cols - len(row_values))
        for cell_index, value in enumerate(padded_values):
            cell = table.cell(row_index, cell_index)
            cell.text = value
            if row_index == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def _is_separator_row(cells: list[str]) -> bool:
    if not cells:
        return False
    normalized = [cell.replace(" ", "") for cell in cells]
    return all(re.fullmatch(r":?-{3,}:?", cell) for cell in normalized if cell) and all(normalized)


def _add_image_or_fallback(document: Document, text: str) -> None:
    match = IMAGE_RE.fullmatch(text.strip())
    if not match:
        document.add_paragraph(text)
        return
    alt = match.group("alt") or "图片"
    image_path = Path(match.group("path"))
    if image_path.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Cm(15.5))
        caption = document.add_paragraph(alt)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"[图片加载失败：原路径为 {image_path.as_posix()}]")
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
