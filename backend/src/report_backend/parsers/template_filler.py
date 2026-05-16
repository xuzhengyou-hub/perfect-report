from __future__ import annotations

import json
import re
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.shared import Cm
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run


PLACEHOLDER_PREFIX = "待用户填写"


@dataclass(slots=True)
class TemplateReport:
    cover_fields: dict[str, str | None]
    roster_fields: dict[str, str | None]
    evaluation_fields: dict[str, str | None]
    sections: dict[str, list[str]]
    images: dict[str, str | None]
    gallery_images: list[str]


@dataclass(slots=True)
class ParagraphTemplate:
    paragraph_properties_xml: Any | None
    run_properties_xml: Any | None


def render_template_report(template_path: Path, output_path: Path, report: TemplateReport) -> Path:
    document = Document(template_path)

    _fill_cover_tables(document, report)
    _fill_experiment_name_paragraph(document, report.cover_fields.get("experiment_name"))
    _fill_body_sections(document, report)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def render_report_preview_markdown(report: TemplateReport, title: str) -> str:
    sections = [f"# {title}"]

    cover_lines = [
        f"- {label}: {report.cover_fields.get(key) or _placeholder(label)}"
        for key, label in (
            ("course_name", "课程名称"),
            ("experiment_name", "实验项目"),
            ("experiment_time", "实验时间"),
            ("class_name", "实验班级"),
            ("teacher_name", "指导教师"),
        )
    ]
    roster_lines = [
        f"- {label}: {report.roster_fields.get(key) or _placeholder(label)}"
        for key, label in (
            ("college", "学院"),
            ("major", "专业"),
            ("student_name", "姓名"),
            ("student_id", "学号"),
        )
    ]

    sections.append("## 模板字段")
    sections.extend(cover_lines)
    sections.extend(roster_lines)

    heading_map = {
        "objective": "实验目的",
        "content": "实验内容",
        "steps": "实验步骤",
        "sequence_components": "序列图组成",
        "collaboration_components": "协作图组成",
        "event_flow": "分析飞机订票的事件流",
        "sequence_idea": "创建飞机订票系统序列图的基本思路",
        "collaboration_steps": "创建飞机订票系统协作图的操作步骤",
        "summary": "实验小结",
    }
    for key, heading in heading_map.items():
        sections.append(f"## {heading}")
        sections.extend(report.sections.get(key) or [_placeholder(heading)])

    image_lines = []
    for key, label in (("sequence_diagram", "序列图"), ("collaboration_diagram", "协作图")):
        image_path = report.images.get(key)
        image_lines.append(f"- {label}: {image_path or _placeholder(label)}")
    for image_path in report.gallery_images:
        if image_path not in report.images.values():
            image_lines.append(f"- 附图: {image_path}")
    sections.append("## 图片槽位")
    sections.extend(image_lines)

    return "\n\n".join(part for part in sections if part).strip()


def parse_template_report(payload: str) -> TemplateReport:
    data = json.loads(payload)
    return TemplateReport(
        cover_fields=_coerce_map(data.get("cover_fields")),
        roster_fields=_coerce_map(data.get("roster_fields")),
        evaluation_fields=_coerce_map(data.get("evaluation_fields")),
        sections=_coerce_sections(data.get("sections")),
        images=_coerce_map(data.get("images")),
        gallery_images=_coerce_list(data.get("gallery_images")),
    )


def finalize_template_report(
    report: TemplateReport,
    *,
    synthesized_markdown: str,
    parsed_markdown: str,
    image_refs: list[str],
) -> TemplateReport:
    detailed_steps = _extract_detailed_steps_from_markdown(synthesized_markdown or parsed_markdown)
    if detailed_steps:
        report.sections["steps"] = _prefer_detailed_steps(detailed_steps, report.sections.get("steps", []))

    normalized_refs = [_canonical_image_path(path) for path in image_refs]
    named_images = {key: _canonical_image_path(value) if value else None for key, value in report.images.items()}
    selected = _select_report_images(normalized_refs)
    if not named_images.get("sequence_diagram"):
        named_images["sequence_diagram"] = selected[0]
    if not named_images.get("collaboration_diagram"):
        named_images["collaboration_diagram"] = selected[1]
    report.images = named_images
    report.gallery_images = _unique_paths(
        [*named_images.values(), *[_canonical_image_path(path) for path in report.gallery_images], *normalized_refs]
    )
    return report


def build_fallback_template_report(
    *,
    title: str,
    extra_info: str,
    synthesized_context: str,
    parsed_markdown: str,
    image_refs: list[str],
) -> TemplateReport:
    source_text = "\n".join(part for part in [extra_info, synthesized_context, parsed_markdown] if part).strip()
    lines = _clean_source_lines(source_text)
    numbered_lines = _extract_numbered_lines(lines)

    sequence_image, collaboration_image = _select_report_images(image_refs)

    objective = _first_matching_lines(lines, ("目的", "目标", "掌握", "理解"), limit=3)
    if not objective:
        objective = [
            f"理解并完成《{title}》相关实验内容。",
            "根据提供资料整理关键流程、图示和结论。",
        ]

    content = _first_matching_lines(lines, ("内容", "包括", "围绕", "系统"), limit=4)
    if not content:
        content = [
            "结合模板要求整理实验背景、关键步骤和分析结果。",
            "对资料中的图示与文字说明进行归纳并写回模板指定位置。",
        ]

    steps = numbered_lines[:6]
    if not steps:
        steps = [
            "整理提供的实验资料与图示。",
            "提取可确认的实验信息并填入模板。",
            "对无法确认的字段保留待用户填写标记。",
        ]

    event_flow = numbered_lines[:8] or steps
    summary = _first_matching_lines(lines, ("小结", "总结", "结论"), limit=3)
    if not summary:
        summary = [
            "已依据现有资料完成模板化填充。",
            "未能从资料中确认的字段已保留为待用户填写。",
        ]

    return TemplateReport(
        cover_fields={
            "course_name": _match_single_value(lines, ("课程名称", "课程", "软件工程")),
            "experiment_name": title,
            "experiment_time": _match_date(source_text),
            "class_name": None,
            "total_score": None,
            "teacher_name": _match_single_value(lines, ("指导教师", "教师", "黄超")),
        },
        roster_fields={
            "college": _match_single_value(lines, ("学院", "计算机科学学院")),
            "major": _match_single_value(lines, ("专业", "计算机科学与技术")),
            "class_name": None,
            "score": None,
            "student_name": None,
            "student_id": None,
            "group_name": None,
            "group_members": None,
            "location": None,
            "experiment_date": _match_date(source_text),
            "teacher_signature": None,
        },
        evaluation_fields={
            "preview_status": None,
            "operation_status": None,
            "attendance_status": None,
            "data_processing_status": None,
        },
        sections={
            "objective": objective,
            "content": content,
            "steps": steps,
            "sequence_components": _first_matching_lines(lines, ("序列图", "时序图", "lifeline", "message"), limit=4),
            "collaboration_components": _first_matching_lines(lines, ("协作图", "通信图", "link", "对象"), limit=4),
            "event_flow": event_flow,
            "sequence_idea": [],
            "collaboration_steps": [],
            "summary": summary,
        },
        images={
            "sequence_diagram": sequence_image,
            "collaboration_diagram": collaboration_image,
        },
        gallery_images=_unique_paths(image_refs),
    )


def _fill_cover_tables(document: DocumentObject, report: TemplateReport) -> None:
    if len(document.tables) < 3:
        raise ValueError("Template does not contain the expected table structure.")

    cover = document.tables[0]
    roster = document.tables[1]
    evaluation = document.tables[2]

    _set_cell_text(cover, 0, 1, report.cover_fields.get("course_name"), "课程名称")
    _set_cell_text(cover, 1, 1, report.cover_fields.get("experiment_name"), "实验项目")
    _set_cell_text(cover, 2, 1, report.cover_fields.get("experiment_time"), "实验时间")
    _set_cell_text(cover, 3, 1, report.cover_fields.get("class_name"), "实验班级")
    _set_cell_text(cover, 4, 1, report.cover_fields.get("total_score"), "总份数")
    _set_cell_text(cover, 5, 1, report.cover_fields.get("teacher_name"), "指导教师")

    _set_cell_text(roster, 0, 1, report.roster_fields.get("college"), "学院")
    _set_cell_text(roster, 0, 6, report.roster_fields.get("major"), "专业")
    _set_cell_text(roster, 0, 11, report.roster_fields.get("class_name"), "班级")
    _set_cell_text(roster, 0, 14, report.roster_fields.get("score"), "成绩")

    _set_cell_text(roster, 1, 1, report.roster_fields.get("student_name"), "姓名")
    _set_cell_text(roster, 1, 5, report.roster_fields.get("student_id"), "学号")
    _set_cell_text(roster, 1, 10, report.roster_fields.get("group_name"), "组别")
    _set_cell_text(roster, 1, 13, report.roster_fields.get("group_members"), "组员")

    _set_cell_text(roster, 2, 2, report.roster_fields.get("location"), "实验地点")
    _set_cell_text(roster, 2, 7, report.roster_fields.get("experiment_date"), "实验日期")
    _set_cell_text(roster, 2, 14, report.roster_fields.get("teacher_signature"), "指导教师签名")

    _set_cell_text(evaluation, 0, 1, report.evaluation_fields.get("preview_status"), "预习情况")
    _set_cell_text(evaluation, 0, 3, report.evaluation_fields.get("operation_status"), "操作情况")
    _set_cell_text(evaluation, 0, 5, report.evaluation_fields.get("attendance_status"), "考勤情况")
    _set_cell_text(evaluation, 0, 7, report.evaluation_fields.get("data_processing_status"), "数据处理情况")


def _fill_experiment_name_paragraph(document: DocumentObject, experiment_name: str | None) -> None:
    paragraph = _find_paragraph(document, "项目名称")
    if paragraph is None:
        return
    original = paragraph.text or "实验       项目名称："
    prefix, separator, _ = original.partition("：")
    if not separator:
        prefix, separator, _ = original.partition(":")
        separator = separator or "："
    _write_paragraph_text(
        paragraph,
        f"{prefix}{separator} {_resolved_value(experiment_name, '项目名称', keep_existing=False)}",
    )


def _fill_body_sections(document: DocumentObject, report: TemplateReport) -> None:
    if _fill_compact_body_sections(document, report):
        return

    section_order = [
        ("objective", "实验目的"),
        ("content", "实验内容"),
        ("steps", "实验步骤"),
        ("sequence_components", "序列图组成"),
        ("collaboration_components", "协作图组成"),
        ("event_flow", "事件流"),
        ("sequence_idea", "序列图的基本思路"),
        ("collaboration_steps", "协作图的操作步骤"),
        ("summary", "实验小结"),
    ]
    anchors: list[tuple[str, Paragraph]] = []
    for key, keyword in section_order:
        paragraph = _find_paragraph(document, keyword)
        if paragraph is not None:
            anchors.append((key, paragraph))

    for index in range(len(anchors) - 1, -1, -1):
        key, anchor = anchors[index]
        next_anchor = anchors[index + 1][1] if index + 1 < len(anchors) else None
        template = _capture_body_template(anchor, next_anchor)
        _clear_between(anchor, next_anchor)
        blocks = _build_section_blocks(key, report)
        _insert_blocks_after(anchor, next_anchor, blocks, template)


def _fill_compact_body_sections(document: DocumentObject, report: TemplateReport) -> bool:
    compact_order = [
        ("purpose_bundle", "实验目的和要求"),
        ("content_bundle", "实验内容及步骤"),
        ("result_bundle", "实验结果"),
        ("discussion_bundle", "问题与讨论"),
    ]
    anchors: list[tuple[str, Paragraph]] = []
    for key, keyword in compact_order:
        paragraph = _find_paragraph(document, keyword)
        if paragraph is not None:
            anchors.append((key, paragraph))
    if len(anchors) < 2:
        return False

    for index in range(len(anchors) - 1, -1, -1):
        key, anchor = anchors[index]
        next_anchor = anchors[index + 1][1] if index + 1 < len(anchors) else None
        template = _capture_body_template(anchor, next_anchor)
        _clear_between(anchor, next_anchor)
        blocks = _build_compact_section_blocks(key, report)
        _insert_blocks_after(anchor, next_anchor, blocks, template)
    return True


def _build_section_blocks(key: str, report: TemplateReport) -> list[dict[str, str]]:
    lines = report.sections.get(key) or []
    blocks: list[dict[str, str]] = [{"type": "paragraph", "value": line} for line in lines if line.strip()]

    if key == "sequence_idea":
        image_path = report.images.get("sequence_diagram")
        if image_path:
            blocks.append({"type": "image", "value": image_path})
    elif key == "collaboration_steps":
        image_path = report.images.get("collaboration_diagram")
        if image_path:
            blocks.append({"type": "image", "value": image_path})

    if not blocks:
        blocks.append({"type": "paragraph", "value": _placeholder(key)})
    return blocks


def _build_compact_section_blocks(key: str, report: TemplateReport) -> list[dict[str, str]]:
    blocks: list[dict[str, str]] = []
    if key == "purpose_bundle":
        blocks.extend(_labeled_paragraph_blocks("实验目的", report.sections.get("objective", [])))
    elif key == "content_bundle":
        blocks.extend(_labeled_paragraph_blocks("实验内容", report.sections.get("content", [])))
        blocks.extend(_labeled_paragraph_blocks("实验步骤", report.sections.get("steps", [])))
    elif key == "result_bundle":
        summary_lines = report.sections.get("summary", [])
        if not summary_lines:
            summary_lines = report.sections.get("event_flow", [])
        blocks.extend(_labeled_paragraph_blocks("实验结果与分析", summary_lines))
        for image_path in report.gallery_images:
            blocks.append({"type": "image", "value": image_path})
    elif key == "discussion_bundle":
        lines = report.sections.get("event_flow", []) or report.sections.get("summary", [])
        blocks.extend(_labeled_paragraph_blocks("补充说明", lines))

    if not blocks:
        blocks.append({"type": "paragraph", "value": _placeholder(key)})
    return blocks


def _labeled_paragraph_blocks(label: str, lines: list[str]) -> list[dict[str, str]]:
    cleaned = [line.strip() for line in lines if line and line.strip()]
    if not cleaned:
        return []
    blocks: list[dict[str, str]] = [{"type": "paragraph", "value": f"{label}："}]
    blocks.extend({"type": "paragraph", "value": line} for line in cleaned)
    return blocks


def _insert_blocks_after(
    anchor: Paragraph,
    next_anchor: Paragraph | None,
    blocks: list[dict[str, str]],
    template: ParagraphTemplate | None,
) -> None:
    sentinel = _insert_paragraph_after(anchor) if next_anchor is None else None
    insertion_point = anchor
    for block in blocks:
        paragraph = next_anchor.insert_paragraph_before() if next_anchor is not None else sentinel.insert_paragraph_before()
        _apply_paragraph_template(paragraph, template)
        if block["type"] == "image":
            _clear_paragraph_content(paragraph)
            run = paragraph.add_run()
            _apply_run_template(run, template)
            run.add_picture(block["value"], width=Cm(13.5))
        else:
            _write_paragraph_text(paragraph, block["value"], template)
        insertion_point = paragraph
    if sentinel is not None:
        sentinel._element.getparent().remove(sentinel._element)


def _find_paragraph(document: DocumentObject, keyword: str) -> Paragraph | None:
    normalized_keyword = _normalize(keyword)
    for paragraph in document.paragraphs:
        if normalized_keyword in _normalize(paragraph.text):
            return paragraph
    return None


def _capture_body_template(anchor: Paragraph, next_anchor: Paragraph | None) -> ParagraphTemplate | None:
    current = anchor._p.getnext()
    stop = next_anchor._p if next_anchor is not None else None

    fallback: Paragraph | None = None
    while current is not None and current is not stop:
        if current.tag.endswith("sectPr"):
            break
        if current.tag.endswith("}p"):
            paragraph = Paragraph(current, anchor._parent)
            if fallback is None:
                fallback = paragraph
            if paragraph.text.strip():
                return _capture_paragraph_template(paragraph)
        current = current.getnext()

    if fallback is not None:
        return _capture_paragraph_template(fallback)
    return _capture_paragraph_template(anchor)


def _capture_paragraph_template(paragraph: Paragraph) -> ParagraphTemplate | None:
    paragraph_properties = deepcopy(paragraph._p.pPr) if paragraph._p.pPr is not None else None
    run_template = _first_run(paragraph)
    run_properties = deepcopy(run_template._r.rPr) if run_template is not None and run_template._r.rPr is not None else None
    return ParagraphTemplate(paragraph_properties_xml=paragraph_properties, run_properties_xml=run_properties)


def _clear_between(anchor: Paragraph, next_anchor: Paragraph | None) -> None:
    current = anchor._p.getnext()
    stop = next_anchor._p if next_anchor is not None else None
    while current is not None and current is not stop:
        if current.tag.endswith("sectPr"):
            break
        next_element = current.getnext()
        current.getparent().remove(current)
        current = next_element


def _insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _set_cell_text(table: Table, row_index: int, column_index: int, value: str | None, label: str) -> None:
    cell = table.cell(row_index, column_index)
    text = _resolved_value(value, label, existing=_normalized_cell_text(cell))
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    template = _capture_paragraph_template(paragraph)
    _write_paragraph_text(paragraph, text, template)
    for extra_paragraph in cell.paragraphs[1:]:
        _clear_paragraph_content(extra_paragraph)


def _write_paragraph_text(paragraph: Paragraph, text: str, template: ParagraphTemplate | None = None) -> None:
    template = template or _capture_paragraph_template(paragraph)
    _apply_paragraph_template(paragraph, template)
    _clear_paragraph_content(paragraph)
    run = paragraph.add_run(text)
    _apply_run_template(run, template)


def _apply_paragraph_template(paragraph: Paragraph, template: ParagraphTemplate | None) -> None:
    if template is None or template.paragraph_properties_xml is None:
        return
    existing = paragraph._p.pPr
    if existing is not None:
        paragraph._p.remove(existing)
    paragraph._p.insert(0, deepcopy(template.paragraph_properties_xml))


def _apply_run_template(run: Run, template: ParagraphTemplate | None) -> None:
    if template is None or template.run_properties_xml is None:
        return
    if run._r.rPr is not None:
        run._r.remove(run._r.rPr)
    run._r.insert(0, deepcopy(template.run_properties_xml))


def _clear_paragraph_content(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag.endswith("}pPr"):
            continue
        paragraph._p.remove(child)


def _first_run(paragraph: Paragraph) -> Run | None:
    for run in paragraph.runs:
        if run.text.strip():
            return run
    return paragraph.runs[0] if paragraph.runs else None


def _resolved_value(value: str | None, label: str, *, existing: str = "", keep_existing: bool = True) -> str:
    normalized_value = (value or "").strip()
    if normalized_value:
        return normalized_value
    if keep_existing and existing:
        return existing
    return _placeholder(label)


def _normalized_cell_text(cell: _Cell) -> str:
    return cell.text.strip()


def _placeholder(label: str) -> str:
    return f"【{PLACEHOLDER_PREFIX}：{label}】"


def _coerce_map(value: object) -> dict[str, str | None]:
    if not isinstance(value, dict):
        return {}
    coerced: dict[str, str | None] = {}
    for key, raw in value.items():
        if raw is None:
            coerced[str(key)] = None
        else:
            text = str(raw).strip()
            coerced[str(key)] = text or None
    return coerced


def _coerce_list(value: object) -> list[str]:
    if not isinstance(value, list):
        return []
    return [str(item).strip() for item in value if str(item).strip()]


def _coerce_sections(value: object) -> dict[str, list[str]]:
    if not isinstance(value, dict):
        return {}
    sections: dict[str, list[str]] = {}
    for key, raw in value.items():
        if isinstance(raw, list):
            sections[str(key)] = [str(item).strip() for item in raw if str(item).strip()]
        elif raw is None:
            sections[str(key)] = []
        else:
            text = str(raw).strip()
            sections[str(key)] = [text] if text else []
    return sections


def _clean_source_lines(source_text: str) -> list[str]:
    cleaned: list[str] = []
    for raw_line in source_text.splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or line.startswith("!["):
            continue
        line = re.sub(r"\*\*", "", line)
        line = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", line)
        cleaned.append(line)
    return cleaned


def _extract_numbered_lines(lines: list[str]) -> list[str]:
    numbered = [line for line in lines if re.match(r"^(\d+[\.\)]|[-*])\s*", line)]
    return [re.sub(r"^(\d+[\.\)]|[-*])\s*", "", line).strip() for line in numbered if line.strip()]


def _extract_detailed_steps_from_markdown(markdown: str) -> list[str]:
    capture = False
    current_heading_level = 0
    parent_step = ""
    steps: list[str] = []

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            continue

        heading_match = re.match(r"^(#{1,6})\s*(.+)$", stripped)
        if heading_match:
            heading_level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            if capture and heading_level <= current_heading_level:
                break
            if "实验内容" in heading_text and "步骤" in heading_text:
                capture = True
                current_heading_level = heading_level
                parent_step = ""
            continue

        if not capture:
            continue

        numbered_match = re.match(r"^\d+[\.\)]\s*(.+)$", stripped)
        if numbered_match:
            parent_step = _strip_markdown_markup(numbered_match.group(1))
            parent_step = parent_step.rstrip("：:")
            if parent_step:
                steps.append(parent_step)
            continue

        bullet_match = re.match(r"^[-*]\s*(.+)$", stripped)
        if bullet_match:
            child = _strip_markdown_markup(bullet_match.group(1))
            if not child:
                continue
            if parent_step:
                combined = f"{parent_step}：{child}"
                if combined not in steps:
                    steps.append(combined)
            else:
                steps.append(child)

    return steps


def _strip_markdown_markup(text: str) -> str:
    cleaned = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    cleaned = re.sub(r"`([^`]+)`", r"\1", cleaned)
    return cleaned.strip()


def _prefer_detailed_steps(detailed_steps: list[str], existing_steps: list[str]) -> list[str]:
    if not detailed_steps:
        return existing_steps
    merged = _unique_strings([*detailed_steps, *existing_steps])
    return merged if len(merged) >= len(existing_steps) else existing_steps


def _unique_strings(items: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for item in items:
        normalized = item.strip()
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        result.append(normalized)
    return result


def _first_matching_lines(lines: list[str], keywords: tuple[str, ...], *, limit: int) -> list[str]:
    matches: list[str] = []
    for line in lines:
        normalized = line.lower()
        if any(keyword.lower() in normalized for keyword in keywords):
            matches.append(line)
        if len(matches) >= limit:
            break
    return matches


def _match_single_value(lines: list[str], keywords: tuple[str, ...]) -> str | None:
    for line in lines:
        normalized = line.lower()
        for keyword in keywords:
            if keyword.lower() in normalized:
                return line
    return None


def _match_date(source_text: str) -> str | None:
    match = re.search(r"(20\d{2}[/-]\d{1,2}[/-]\d{1,2})", source_text)
    if match:
        return match.group(1).replace("-", "/")
    return None


def _select_report_images(image_refs: list[str]) -> tuple[str | None, str | None]:
    sequence_image: str | None = None
    collaboration_image: str | None = None

    for image_ref in image_refs:
        normalized = _normalize(Path(image_ref).stem)
        if sequence_image is None and ("序列" in normalized or "sequence" in normalized or "代码" in normalized):
            sequence_image = image_ref
            continue
        if collaboration_image is None and (
            "协作" in normalized or "communication" in normalized or "collaboration" in normalized or "结果" in normalized
        ):
            collaboration_image = image_ref

    if sequence_image is None and image_refs:
        sequence_image = image_refs[0]
    if collaboration_image is None and len(image_refs) > 1:
        collaboration_image = image_refs[1]
    elif collaboration_image is None:
        collaboration_image = sequence_image

    return sequence_image, collaboration_image


def _unique_paths(paths: list[str | None]) -> list[str]:
    result: list[str] = []
    seen: set[str] = set()
    for path in paths:
        if not path:
            continue
        normalized = _canonical_image_path(path)
        dedupe_key = str(Path(normalized)).lower()
        if not normalized or dedupe_key in seen:
            continue
        seen.add(dedupe_key)
        result.append(normalized)
    return result


def _normalize(text: str) -> str:
    lowered = text.lower()
    return re.sub(r"[\s:：()（）_\-.,，。\\/]+", "", lowered)


def _canonical_image_path(path: str) -> str:
    normalized = str(path).strip()
    if not normalized:
        return normalized
    candidate = Path(normalized)
    try:
        return str(candidate.resolve(strict=False))
    except OSError:
        return normalized
